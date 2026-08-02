#!/usr/bin/env python3
"""Rebuild the internal pocket report with evidence-aware language and layout."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F4F6"
BORDER = "DADCE0"
MUTED = RGBColor(85, 85, 85)

DOCKING_CANDIDATES = """\
MCULE-3294307948\t-11.575\t-8.481\t3.094
MCULE-8779851442\t-7.788\t-4.880\t2.908
romidepsin-active-thiol\t-6.232\t-3.560\t2.672
SEL-MOD-5c5395b9f734\t-11.371\t-8.717\t2.654
MCULE-5058868643\t-10.976\t-8.374\t2.602
MCULE-7946764063\t-10.149\t-7.562\t2.587
MCULE-8397255882\t-10.643\t-8.098\t2.545
METTL7-RESCUE-c1c40a919201\t-10.226\t-7.754\t2.472
METTL7-RESCUE-87c82a1f3ebd\t-8.849\t-6.392\t2.457
CHEMBL5748724\t-7.826\t-5.460\t2.366
CHEMBL5945752\t-8.013\t-5.681\t2.332
T6770\t-8.486\t-6.191\t2.295
MCULE-5935081243\t-8.333\t-6.080\t2.253
CHEMBL5860275\t-7.631\t-5.493\t2.138
CHEMBL5171494\t-8.720\t-6.604\t2.116
T4071\t-7.001\t-4.910\t2.091
SEL-MOD-b4c25a2b1ac9\t-11.529\t-9.485\t2.044
MCULE-3461284601\t-6.574\t-4.553\t2.021
METTL7-RESCUE-84c3290a2032\t-8.482\t-6.516\t1.966
MCULE-7994089601\t-9.042\t-7.081\t1.961
MCULE-4790369949\t-7.911\t-6.011\t1.900
METTL7-RESCUE-54af709eba1e\t-11.023\t-9.138\t1.885
METTL7-RESCUE-5c5867044749\t-8.120\t-6.274\t1.846
MCULE-8950822836\t-8.330\t-6.491\t1.839
MCULE-2420078993\t-8.918\t-7.079\t1.839
METTL7-BRICS-0003\t-9.064\t-7.236\t1.828
MCULE-3188917096\t-11.542\t-9.721\t1.821
CHEMBL5743786\t-7.096\t-5.283\t1.813
MCULE-5545626239\t-9.261\t-7.454\t1.807
T13194\t-6.506\t-4.707\t1.799
MCULE-7154297007\t-10.541\t-8.743\t1.798
CHEMBL6041327\t-7.287\t-5.543\t1.744
MCULE-7080563723\t-10.802\t-9.082\t1.720
MCULE-8355279029\t-9.376\t-7.676\t1.700
MCULE-4515095248\t-10.239\t-8.553\t1.686
CHEMBL5646588\t-9.365\t-7.693\t1.672
MCULE-9365544165\t-10.909\t-9.238\t1.671
MCULE-4988716858\t-10.340\t-8.684\t1.656
MCULE-3281519817\t-9.640\t-7.988\t1.652
METTL7-BRICS-0017\t-8.546\t-6.909\t1.637
MCULE-4227481601\t-6.426\t-4.808\t1.618
"""


def set_run_font(run, name: str, size: float, *, bold: bool | None = None,
                 color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(element, fill: str) -> None:
    props = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_cell_margins(cell, twips: int = 72) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if style:
        result.style = style
    if text:
        result.add_run(text)
    return result


def insert_paragraph_before(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if style:
        result.style = style
    if text:
        result.add_run(text)
    return result


def move_table_after(table, paragraph: Paragraph) -> None:
    paragraph._p.addnext(table._tbl)


def insert_paragraph_after_table(table, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    result = Paragraph(new_p, table._parent)
    if style:
        result.style = style
    if text:
        result.add_run(text)
    return result


def format_table(table, widths: tuple[float, ...] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for column, width in zip(table.columns, widths):
            column.width = Inches(width)
        grid = table._tbl.tblGrid
        for grid_col, width in zip(grid.gridCol_lst, widths):
            grid_col.set(qn("w:w"), str(int(width * 1440)))
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if widths:
                cell.width = Inches(widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if r_idx == 0:
                shade(cell._tc, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, "Arial", 8.5, bold=(r_idx == 0))


def add_metadata_table_after(doc: Document, paragraph: Paragraph) -> None:
    table = doc.add_table(rows=5, cols=2)
    rows = [
        ("Target", "METTL7B · UniProt Q6UX53"),
        ("Structure", "AlphaFold AF-Q6UX53-F1-model_v6 · chain A · model 1"),
        ("Pocket", "FPOCKET source pocket 2 · internal pocket ID 1"),
        ("Pocket metrics", "47 residues · volume 1,690.5 Å³ · druggability 0.832"),
        ("Evidence status", "Internal computational analysis; no experimental validation is claimed"),
    ]
    for row, values in zip(table.rows, rows):
        row.cells[0].text, row.cells[1].text = values
    format_table(table, (1.4, 5.5))
    for row in table.rows:
        shade(row.cells[0]._tc, LIGHT_GRAY)
        row.cells[0].paragraphs[0].runs[0].bold = True
    move_table_after(table, paragraph)


def add_evidence_table_after(doc: Document, paragraph: Paragraph) -> None:
    heading = insert_paragraph_after(paragraph, "Evidence and interpretation", "Heading 2")
    intro = insert_paragraph_after(
        heading,
        "All evidence in this section was generated by our computational workflow. "
        "The model was explicitly parameterized with a Cys202–Cys203 covalent bond before "
        "production dynamics; the trajectory therefore tests short-timescale structural "
        "accommodation of that imposed oxidized state, not spontaneous bond formation.",
    )
    caption = insert_paragraph_after(
        intro,
        "Table 5. Computational observations and the interpretation supported by each result.",
    )
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(3)
    for run in caption.runs:
        set_run_font(run, "Arial", 8.5)
        run.italic = True
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Computational evidence", "Observed result", "Specific meaning")):
        cell.text = text
    rows = [
        (
            "Pre-bonded explicit-solvent OpenMM trajectory",
            "200 ps production after staged accommodation and equilibration",
            "The imposed oxidized topology did not undergo immediate local structural failure over the sampled interval.",
        ),
        (
            "Cys202–Cys203 SG distance",
            "2.03 ± 0.05 Å; range 1.95–2.16 Å",
            "The covalently parameterized bond retained expected disulfide geometry. Because the bond was imposed, this does not measure formation propensity.",
        ),
        (
            "Sulfur solvent accessibility",
            "Cys202 mean SG SASA ≈ 0.040 nm²; Cys148 ≈ 0 nm²; Cys203 lower than Cys202",
            "In this oxidized trajectory, Cys202 remained partly exposed while Cys148 was buried. This does not establish accessibility in the reduced state.",
        ),
        (
            "Local Cα RMSF",
            "Cys202 0.482 Å; Cys203 0.422 Å; residues 195–205 span 0.366–0.642 Å",
            "The bonded cysteine pair was not unusually mobile within this single oxidized trajectory. A reduced control is required to claim oxidation-induced rigidification.",
        ),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            cell.text = text
    format_table(table, (2.05, 1.65, 3.2))
    move_table_after(table, caption)
    last = insert_paragraph_after_table(
        table,
        "Supported conclusion: a pre-bonded Cys202–Cys203 oxidized model is mechanically "
        "accommodated over this short simulation. Not established: spontaneous disulfide "
        "formation, thermodynamic favorability, physiological oxidation, or a biological redox-switch function.",
    )
    last.paragraph_format.space_before = Pt(3)
    last.paragraph_format.space_after = Pt(3)
    for run in last.runs:
        run.italic = True


def add_story_figure_after(doc: Document, paragraph: Paragraph, image_path: Path) -> None:
    figure = insert_paragraph_after(paragraph)
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_before = Pt(8)
    figure.paragraph_format.space_after = Pt(4)
    figure.add_run().add_picture(str(image_path), width=Inches(7.0))
    caption = insert_paragraph_after(
        figure,
        "Figure 1. Integrated computational view of the METTL7B sequence. Lime cells identify "
        "the 47 residues assigned to FPOCKET pocket 2; the blue top edge reports docking-contact "
        "frequency for ligands scoring < −5 in the selected run; the amber edge reports ESMC "
        "sequence constraint. The layers answer different questions and should be interpreted "
        "together rather than as independent pocket calls.",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    for run in caption.runs:
        set_run_font(run, "Arial", 8.5, color="555555")
        run.italic = True


def add_docking_candidates_after(doc: Document, paragraph: Paragraph) -> None:
    heading = insert_paragraph_after(paragraph, "9. Differential docking candidates", "Heading 1")
    intro = insert_paragraph_after(
        heading,
        "Table 4 ranks computational docking candidates by the score difference between the "
        "7A and 7B runs. Δ score is defined as 7A − 7B; positive values therefore indicate a "
        "more favorable (more negative) score against 7B. Docking scores are ranking evidence, "
        "not measured binding affinities.",
    )
    caption = insert_paragraph_after(
        intro,
        "Table 4. Candidates with the largest positive differential docking scores for 7B relative to 7A.",
    )
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(3)
    for run in caption.runs:
        set_run_font(run, "Arial", 8.5)
        run.italic = True

    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Candidate"
    table.rows[0].cells[1].text = "7B score"
    table.rows[0].cells[2].text = "7A score"
    table.rows[0].cells[3].text = "Δ (7A − 7B)"
    for line in DOCKING_CANDIDATES.strip().splitlines():
        candidate, score_7b, score_7a, delta = line.split("\t")
        cells = table.add_row().cells
        for cell, value in zip(cells, (candidate, score_7b, score_7a, delta)):
            cell.text = value
        for cell in cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    format_table(table, (3.6, 1.05, 1.05, 1.2))
    move_table_after(table, caption)
    after = insert_paragraph_after_table(
        table,
        "The UUID and internal run/pose identifiers are intentionally omitted from the manuscript "
        "table; they remain database provenance rather than interpretive result columns.",
    )
    after.paragraph_format.space_before = Pt(3)
    for run in after.runs:
        set_run_font(run, "Arial", 8.5, color="555555")
        run.italic = True


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name, size, color, before, after in (
        ("Heading 1", 13, BLUE, 8, 3),
        ("Heading 2", 11.5, BLUE, 8, 3),
        ("Heading 3", 10.5, DARK_BLUE, 6, 2),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Report Title" not in styles:
        title = styles.add_style("Report Title", 1)
    else:
        title = styles["Report Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title.font.size = Pt(18)
    title.font.bold = False
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(7.5)
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.keep_together = False
    code.paragraph_format.keep_with_next = False


def replace_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def revise(source: Path, output: Path, image_path: Path) -> None:
    doc = Document(source)
    original = list(doc.paragraphs)
    configure_styles(doc)

    # Replace the undersized legacy screenshot with the full-width story figure.
    for shape in list(doc.inline_shapes):
        inline = shape._inline
        inline.getparent().remove(inline)

    for section in doc.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    replacements = {
        19: "3. Pocket residue network",
        21: "4. The Cys148–Cys202 sulfur wall",
        27: "5. Position of Cys203",
        34: "6. Functional organization of the pocket",
        53: "7. Implications for ligand recognition",
        59: "8. Conclusions",
        5: (
            "We computationally characterized the principal METTL7B pocket (UniProt Q6UX53) by "
            "integrating FPOCKET geometry, P2Rank prediction, residue-neighborhood analysis, "
            "ligand-conditioned SAM/SAH docking contacts, docking-run contact frequency, and ESMC "
            "sequence constraint. FPOCKET pocket 2 is a 47-residue, 1,690.5 Å³ cavity with a "
            "druggability score of 0.832. The integrated view identifies a chemically heterogeneous "
            "cavity with a predominantly hydrophobic/aromatic face and a sulfur/polar face. These "
            "results define a consensus computational candidate for ligand recognition and design; "
            "they do not constitute experimental validation."
        ),
        7: (
            "Three computational layers were used for distinct purposes. FPOCKET provides geometric "
            "cavity detection, P2Rank provides an independent structure-based pocket prediction, and "
            "BioHub supplies ligand-conditioned docking/contact evidence for SAM and SAH. BioHub is "
            "therefore not treated as a third pocket detector."
        ),
        9: (
            "FPOCKET and P2Rank define related but non-identical cavity boundaries. SAM/SAH contact "
            "residues from BioHub and docking-run contact frequencies are overlaid on those structural "
            "predictions to show which parts of the cavity are actually sampled by ligands."
        ),
        10: (
            "Docked SAM and SAH occupy the same broad cavity as the selected structural pocket. "
            "Residues contacted by these ligand-conditioned calculations support the selected pocket "
            "and can also identify adjacent residues outside the FPOCKET boundary that warrant inspection."
        ),
        12: "FPOCKET pocket 2 comprises the following 47 residues:",
        16: (
            "One face is predominantly hydrophobic and aromatic, with Phe36, Met40, and Phe103 "
            "providing nonpolar surface. Gly80 helps shape the cavity geometry, whereas Asp98 marks "
            "a charged/polar boundary; neither is classified here as a hydrophobic side chain."
        ),
        18: (
            "The opposing wall presents a chemically diverse environment composed of Val146, "
            "Cys148, Ser149, His175, Trp195, Gly199, Asp200, Gly201, and Cys202. This surface "
            "contains the principal sulfur-bearing and polar residues of the selected cavity."
        ),
        22: (
            "A notable feature of the static model is the placement of Cys148 and Cys202 on the same "
            "cavity-facing wall. Their side-chain orientation and separation define a sulfur-rich "
            "microenvironment; static geometry alone is not a quantitative solvent-accessibility measurement."
        ),
        24: (
            "Visualization shows that the Cys148 and Cys202 side chains project toward the selected "
            "cavity. Accessibility claims are reserved for explicit SASA calculations and depend on "
            "the modeled redox state."
        ),
        25: (
            "The selected FPOCKET pocket therefore contains a dual-cysteine wall in the static model, "
            "with Cys202 as a pocket residue and Cys148 as its closest sulfur partner."
        ),
        31: "Cys202 is the principal sulfur-containing residue within the selected FPOCKET pocket.",
        32: "Cys148 is the closest sulfur partner to Cys202 in the static model (SG–SG 3.81 Å).",
        56: (
            "Cys148 and Cys202 define two nearby sulfur-bearing side chains in the static pocket model. "
            "Their ligand accessibility and reactivity require state-specific computational or experimental testing."
        ),
        58: (
            "Agreement among FPOCKET geometry, P2Rank prediction, ligand-conditioned SAM/SAH contacts, "
            "docking-contact frequency, and structural visualization supports this cavity as a consensus "
            "computational candidate for ligand binding."
        ),
        60: (
            "Our integrated computational analysis identifies FPOCKET pocket 2 as a coherent, druggable "
            "METTL7B cavity containing a dense residue network and a distinctive sulfur/polar surface. "
            "The 47-residue sequence map shows how structural membership, docking-contact frequency, and "
            "ESMC constraint overlap at residue resolution. This combined view prioritizes testable regions "
            "for ligand design while keeping geometric prediction, ligand-conditioned evidence, and sequence "
            "constraint conceptually separate."
        ),
        64: "Short-Timescale Structural Accommodation of a Pre-Bonded Cys202–Cys203 Vicinal Disulfide Model in Human METTL7B",
        66: (
            "We used an internal explicit-solvent OpenMM workflow to test whether a covalently "
            "parameterized Cys202–Cys203 vicinal disulfide model could be accommodated by the METTL7B "
            "fold over a short trajectory. After staged sulfur accommodation, minimization, NVT/NPT "
            "equilibration, and 200 ps of production dynamics, the imposed bond retained an SG–SG "
            "distance of 2.03 ± 0.05 Å (range 1.95–2.16 Å). Cys202 remained partly solvent exposed, "
            "and local Cα RMSF values did not indicate immediate disruption around residues 195–205. "
            "These computational results support short-timescale mechanical accommodation of the "
            "pre-bonded oxidized topology. They do not show spontaneous disulfide formation, "
            "thermodynamic favorability, physiological oxidation, or redox-switch function."
        ),
        70: (
            "Human METTL7B contains an adjacent Cys202–Cys203 motif near the computationally selected "
            "ligand-binding pocket. No experimental structure or biochemical measurement is used here "
            "to establish oxidation of this motif; the present work asks a narrower computational question."
        ),
        71: (
            "The objective was to determine whether an explicitly pre-bonded oxidized Cys202–Cys203 "
            "topology is mechanically compatible with the METTL7B fold over a short simulation. This "
            "design tests accommodation after imposing the covalent state, not the probability or pathway "
            "of disulfide formation."
        ),
        85: "The imposed Cys202–Cys203 disulfide geometry remained stable throughout the 200 ps production trajectory.",
        90: (
            "These distances are consistent with expected protein disulfide geometry. Because the "
            "covalent bond was part of the force-field topology, the result shows retention of the imposed "
            "bond geometry rather than spontaneous formation."
        ),
        96: (
            "Within this pre-bonded oxidized trajectory, Cys202 remained partly solvent accessible while "
            "Cys148 was buried. These values are state- and trajectory-specific and should not be generalized "
            "to the reduced model without a matched control."
        ),
        104: (
            "The simulation shows that a pre-bonded Cys202–Cys203 vicinal disulfide can be accommodated "
            "within the METTL7B structure without immediate local instability over 200 ps."
        ),
        105: (
            "The SG–SG distance remained within the expected range for a covalently parameterized "
            "disulfide. This is a consistency check on the oxidized model, not evidence for spontaneous "
            "bond formation or thermodynamic favorability."
        ),
        108: (
            "Overall, our computational results support the short-timescale structural accommodation "
            "of the imposed oxidized topology. They do not establish physiological formation, relative "
            "free energy, biological stability, or a functional redox-switch mechanism."
        ),
        130: (
            "Our explicit-solvent OpenMM trajectory supports a narrow conclusion: a covalently "
            "parameterized Cys202–Cys203 vicinal disulfide model remains mechanically accommodated "
            "during 200 ps of production dynamics. The result provides a reproducible starting point "
            "for longer replicated oxidized simulations and a matched reduced-state control; it is not "
            "evidence that the bond forms spontaneously or functions physiologically."
        ),
    }
    for idx, text in replacements.items():
        replace_text(original[idx], text)

    # Remove legacy paragraph rules that imported as stray horizontal lines.
    for paragraph in original:
        p_pr = paragraph._p.get_or_add_pPr()
        border = p_pr.find(qn("w:pBdr"))
        if border is not None:
            p_pr.remove(border)

    original[0].style = "Report Title"
    original[64].style = "Report Title"
    original[64].paragraph_format.page_break_before = True
    for idx in (4, 15, 17):
        original[idx].style = "Heading 2"

    add_metadata_table_after(doc, original[5])
    add_story_figure_after(doc, original[10], image_path)
    add_docking_candidates_after(doc, original[60])
    add_evidence_table_after(doc, original[76])

    # Keep the compact pocket-property table together instead of splitting it
    # between the first two pages.
    for table in doc.tables:
        if table.cell(0, 0).text.strip() == "Property":
            page_break = OxmlElement("w:p")
            table._tbl.addprevious(page_break)
            table_heading = Paragraph(page_break, table._parent)
            table_heading.paragraph_format.page_break_before = True
            table_heading.paragraph_format.space_after = Pt(3)
            run = table_heading.add_run("Table 1. Structural properties of METTL7B FPOCKET pocket 2.")
            set_run_font(run, "Arial", 8.5)
            run.italic = True
        elif table.cell(0, 0).text.strip() == "Residue":
            caption_p = OxmlElement("w:p")
            table._tbl.addprevious(caption_p)
            caption = Paragraph(caption_p, table._parent)
            caption.paragraph_format.space_after = Pt(3)
            run = caption.add_run("Table 2. Residue-neighborhood network within 6 Å.")
            set_run_font(run, "Arial", 8.5)
            run.italic = True
        elif table.cell(0, 0).text.strip() == "Sulfur Pair":
            caption_p = OxmlElement("w:p")
            table._tbl.addprevious(caption_p)
            caption = Paragraph(caption_p, table._parent)
            caption.paragraph_format.space_after = Pt(3)
            run = caption.add_run("Table 3. Static sulfur-to-sulfur distances in the METTL7B model.")
            set_run_font(run, "Arial", 8.5)
            run.italic = True

    appendix_a = insert_paragraph_before(original[133], "Appendix A. Internal computational data", "Heading 1")
    appendix_a.paragraph_format.page_break_before = True
    insert_paragraph_before(original[133], "A1. Local Cα RMSF values", "Heading 2")
    insert_paragraph_before(
        original[133],
        "Machine-readable values produced by our MDTraj analysis; distances and fluctuations are reported in ångströms.",
    )
    insert_paragraph_before(original[148], "A2. Production-frame sulfur metrics", "Heading 2")
    insert_paragraph_before(
        original[148],
        "Frame-level values from the 200 ps production trajectory; SASA values are reported in nm².",
    )
    appendix_b = insert_paragraph_before(original[174], "Appendix B. Reproducible OpenMM and MDTraj workflow", "Heading 1")
    appendix_b.paragraph_format.page_break_before = True
    insert_paragraph_before(
        original[174],
        "The following code is preserved as computational provenance. Monospace shaded paragraphs are the document equivalent of inline <code> formatting.",
    )

    for idx in list(range(133, 169)) + list(range(174, len(original))):
        paragraph = original[idx]
        paragraph.style = "Code Block"
        shade(paragraph._p, LIGHT_GRAY)
        for run in paragraph.runs:
            set_run_font(run, "Consolas", 7.5)

    for table in doc.tables:
        format_table(table)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=Path("papers/totah.lab.docx"))
    parser.add_argument("--output", type=Path, default=Path("papers/totah.lab.revised.docx"))
    parser.add_argument(
        "--figure",
        type=Path,
        default=Path("/Users/yazan/Desktop/Screenshot 2026-07-30 at 3.43.38\u202fPM.png"),
    )
    args = parser.parse_args()
    revise(args.source.resolve(), args.output.resolve(), args.figure.resolve())


if __name__ == "__main__":
    main()
